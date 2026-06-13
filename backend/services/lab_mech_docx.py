# backend/services/lab_mech_docx.py
from __future__ import annotations

import os
import shutil
import tempfile
import zipfile
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn


# =========================================================
# public functions
# =========================================================
def safe_open_docx(path: str) -> Document:
    """
    安全開啟 docx：
    1. 檢查存在
    2. 檢查大小
    3. 驗證 zip
    4. 複製到 temp 避免鎖檔
    """
    if not os.path.exists(path):
        raise FileNotFoundError(path)

    size = os.path.getsize(path)
    if size < 1024:
        raise RuntimeError(f"docx file too small ({size} bytes): {path}")

    try:
        with zipfile.ZipFile(path, "r") as z:
            _ = z.namelist()
    except Exception as e:
        raise RuntimeError(f"docx is not a valid zip/package: {path} err={e}")

    tmp_dir = tempfile.mkdtemp(prefix="labmech_docx_")
    tmp_path = os.path.join(tmp_dir, os.path.basename(path))
    shutil.copy2(path, tmp_path)

    return Document(tmp_path)


def resolve_mech_template_path(docx_name: str, uploads_base: str) -> Dict[str, Any]:
    docx_name = os.path.basename((docx_name or "").strip())
    docx_path = os.path.join(uploads_base, "templates", docx_name)
    return {
        "docx_name": docx_name,
        "base": uploads_base,
        "docx_path": docx_path,
    }


def fill_header_placeholders(doc: Document, report: Dict[str, Any]):
    mapping = {
        "${report_no}": report.get("report_no"),
        "${entrust_no}": report.get("entrust_no"),
        "${product_name}": report.get("product_name"),
        "${spec_desc}": report.get("spec_desc"),
        "${lot_no}": report.get("lot_no"),
        "${lot_qty}": report.get("lot_qty"),
        "${plating}": report.get("plating"),
        "${material}": report.get("material"),
        "${manufacturer}": report.get("manufacturer"),
        "${standard_type}": report.get("standard_type"),
        "${standard_desc}": report.get("standard_desc"),
        "${env_temp}": report.get("env_temp"),
        "${env_humidity}": report.get("env_humidity"),
        "${test_date}": report.get("test_date"),
        "${complete_date}": report.get("complete_date"),
        "${tester}": report.get("tester"),
        "${reviewer}": report.get("reviewer"),
        "${remarks}": report.get("remarks"),
    }
    _replace_text_everywhere(doc, mapping)


def fill_mech_docx_tables(doc: Document, report: Dict[str, Any]):
    """
    固定表格填值主入口
    """
    _fill_basic_header_tables(doc, report)

    items = report.get("items", []) or []
    items_by_code = {x.get("test_code"): x for x in items}

    for table in doc.tables:
        if items_by_code.get("core_hardness"):
            _fill_hardness_block(table, items_by_code["core_hardness"])
        elif items_by_code.get("surface_hardness"):
            _fill_hardness_block(table, items_by_code["surface_hardness"])

        if items_by_code.get("decarb"):
            _fill_decarb_block(table, items_by_code["decarb"])

        if items_by_code.get("salt_spray"):
            _fill_salt_spray_block(table, items_by_code["salt_spray"])

        if items_by_code.get("hydrogen"):
            _fill_hydrogen_block(table, items_by_code["hydrogen"])


# =========================================================
# internal helpers
# =========================================================
def _apply_default_run_font(run):
    font_name = "DFKai-SB"
    run.font.name = font_name
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
        run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    except Exception:
        pass


def _set_cell_text(cell, text: Any, red: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    _apply_default_run_font(run)
    if red:
        run.font.color.rgb = RGBColor(255, 0, 0)


def _replace_text_everywhere(doc: Document, mapping: Dict[str, Any]):
    def replace_in_paragraph(p):
        full = p.text or ""
        changed = False
        for k, v in mapping.items():
            if k in full:
                full = full.replace(k, "" if v is None else str(v))
                changed = True
        if changed:
            p.clear()
            run = p.add_run(full)
            _apply_default_run_font(run)

    for p in doc.paragraphs:
        replace_in_paragraph(p)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)


def _norm_text(s: Any) -> str:
    return ("" if s is None else str(s)).replace("\n", " ").replace("\r", " ").strip()


def _format_spec_range(spec_min: Any, spec_max: Any) -> str:
    min_text = _norm_text(spec_min)
    max_text = _norm_text(spec_max)
    if min_text and max_text:
        return f"{min_text} ~ {max_text}"
    if min_text:
        return f"{min_text} MIN"
    if max_text:
        return f"{max_text} MAX"
    return ""


def _row_unique_cells(row):
    uniq = []
    seen = set()
    for c in row.cells:
        k = id(c._tc)
        if k in seen:
            continue
        seen.add(k)
        uniq.append(c)
    return uniq


def _find_cell_contains(table, keyword: str):
    keyword = _norm_text(keyword)
    if not keyword:
        return None

    for ri, row in enumerate(table.rows):
        cells = _row_unique_cells(row)
        for ci, cell in enumerate(cells):
            txt = _norm_text(cell.text)
            if keyword in txt:
                return (ri, ci, cell)
    return None


def _find_right_writable_cell(table, ri: int, ci: int, max_scan: int = 8):
    row_cells = _row_unique_cells(table.rows[ri])

    for j in range(ci + 1, min(len(row_cells), ci + 1 + max_scan)):
        txt = _norm_text(row_cells[j].text)
        if txt == "":
            return row_cells[j]

    for j in range(ci + 1, min(len(row_cells), ci + 1 + max_scan)):
        return row_cells[j]

    for rr in range(ri + 1, min(len(table.rows), ri + 3)):
        down_cells = _row_unique_cells(table.rows[rr])
        for cc in range(max(0, ci - 1), min(len(down_cells), ci + 3)):
            txt = _norm_text(down_cells[cc].text)
            if txt == "":
                return down_cells[cc]

    return None


def _set_by_label(table, label: str, value: Any, red: bool = False) -> bool:
    found = _find_cell_contains(table, label)
    if not found:
        return False

    ri, ci, _ = found
    target = _find_right_writable_cell(table, ri, ci)
    if not target:
        return False

    _set_cell_text(target, value, red=red)
    return True


def _find_section_row(table, keyword: str) -> Optional[int]:
    for ri, row in enumerate(table.rows):
        txt = " | ".join(_norm_text(c.text) for c in _row_unique_cells(row))
        if keyword in txt:
            return ri
    return None


def _set_multiline_cell(cell, lines: List[str]):
    cell.text = ""
    p = cell.paragraphs[0]
    for i, line in enumerate(lines):
        run = p.add_run("" if line is None else str(line))
        _apply_default_run_font(run)
        if i < len(lines) - 1:
            run.add_break()


def _fill_basic_header_tables(doc: Document, report: Dict[str, Any]):
    for table in doc.tables:
        _set_by_label(table, "委託編號", report.get("report_no"))
        _set_by_label(table, "品    名", report.get("product_name"))
        _set_by_label(table, "品 名", report.get("product_name"))
        _set_by_label(table, "規    格", report.get("spec_desc"))
        _set_by_label(table, "規 格", report.get("spec_desc"))
        _set_by_label(table, "批    號", report.get("lot_no"))
        _set_by_label(table, "批 號", report.get("lot_no"))
        _set_by_label(table, "批量(pcs)", report.get("lot_qty"))
        _set_by_label(table, "批量", report.get("lot_qty"))
        _set_by_label(table, "鍍    別", report.get("plating"))
        _set_by_label(table, "鍍 別", report.get("plating"))
        _set_by_label(table, "材    質", report.get("material"))
        _set_by_label(table, "材 質", report.get("material"))
        _set_by_label(table, "製造廠商", report.get("manufacturer"))
        _set_by_label(table, "測試日期", report.get("test_date"))
        _set_by_label(table, "測試完成日期", report.get("complete_date"))

        found = _find_cell_contains(table, "檢驗環境")
        if found:
            ri, ci, _ = found
            target = _find_right_writable_cell(table, ri, ci)
            if target:
                env_text = f"溫度：{report.get('env_temp') or ''} ℃；溼度：{report.get('env_humidity') or ''} ％RH"
                _set_cell_text(target, env_text)

        drawing = (report.get("standard_type") or "").strip().lower() == "drawing"
        regulation = (report.get("standard_type") or "").strip().lower() == "regulation"

        found_std = _find_cell_contains(table, "依據標準")
        if found_std:
            ri, _, _ = found_std
            for c in _row_unique_cells(table.rows[ri]):
                t = c.text or ""
                if drawing:
                    t = t.replace("□圖號", "☑圖號")
                if regulation:
                    t = t.replace("□法規", "☑法規")
                if report.get("standard_desc"):
                    if "圖號" in t and drawing and report.get("standard_desc") not in t:
                        t = t.replace("：", f"：{report.get('standard_desc')}", 1) if "：" in t else t
                    if "法規" in t and regulation and report.get("standard_desc") not in t:
                        t = t.replace("：", f"：{report.get('standard_desc')}", 1) if "：" in t else t
                c.text = t


def _fill_hardness_block(table, item: Dict[str, Any]):
    start_ri = _find_section_row(table, "一．硬度測試")
    if start_ri is None:
        return False

    rows = item.get("rows", []) or []
    if not rows:
        return False

    section_row_cells = _row_unique_cells(table.rows[start_ri])
    for c in section_row_cells:
        txt = c.text or ""
        if item.get("test_code") == "core_hardness":
            txt = txt.replace("□心部硬度", "☑心部硬度")
        if item.get("test_code") == "surface_hardness":
            txt = txt.replace("□表面硬度", "☑表面硬度")

        method_code = item.get("method_code") or ""
        if method_code and f"□{method_code}" in txt:
            txt = txt.replace(f"□{method_code}", f"☑{method_code}")
        c.text = txt

    target_cell = None
    for rr in range(start_ri + 1, min(len(table.rows), start_ri + 8)):
        cells = _row_unique_cells(table.rows[rr])
        for c in cells:
            if _norm_text(c.text) == "":
                target_cell = c
                break
        if target_cell:
            break

    if not target_cell:
        return False

    target_cell.text = ""
    p = target_cell.paragraphs[0]

    def add_line(line: str, red: bool = False):
        run = p.add_run(line)
        _apply_default_run_font(run)
        if red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        run.add_break()

    add_line(f"規格：{_format_spec_range(item.get('spec_min'), item.get('spec_max'))}")
    add_line(
        f"合格：{item.get('pass_count') or 0}    不合格：{item.get('fail_count') or 0}    結果：{item.get('result') or ''}",
        red=(item.get("result") == "FAIL"),
    )
    add_line("")

    for r in rows:
        values = r.get("values") or []
        values_text = ", ".join("" if v is None else str(v) for v in values)
        line = (
            f"{r.get('sample_no')}. "
            f"原始值[{values_text}]  "
            f"平均={r.get('avg_value') or ''}  "
            f"判定={r.get('judge_value') or ''}  "
            f"{r.get('result') or ''}"
        )
        add_line(line, red=(r.get("result") == "FAIL"))

    return True


def _fill_decarb_block(table, item: Dict[str, Any]):
    start_ri = _find_section_row(table, "四．脫碳層測試")
    if start_ri is None:
        return False

    rows = item.get("rows", []) or []
    if not rows:
        return False

    target_cell = None
    for rr in range(start_ri + 1, min(len(table.rows), start_ri + 8)):
        cells = _row_unique_cells(table.rows[rr])
        for c in cells:
            if _norm_text(c.text) == "":
                target_cell = c
                break
        if target_cell:
            break

    if not target_cell:
        return False

    target_cell.text = ""
    p = target_cell.paragraphs[0]

    def add_line(line: str, red: bool = False):
        run = p.add_run(line)
        _apply_default_run_font(run)
        if red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        run.add_break()

    add_line("判定公式：HV2 ≥ HV1 - 30 且 HV3 ≤ HV1 + 30")
    add_line(
        f"合格：{item.get('pass_count') or 0}    不合格：{item.get('fail_count') or 0}    結果：{item.get('result') or ''}",
        red=(item.get("result") == "FAIL"),
    )
    add_line("")

    for r in rows:
        line = (
            f"{r.get('sample_no')}. "
            f"HV1={r.get('hv1') or ''}  "
            f"HV2={r.get('hv2') or ''} ({'OK' if r.get('hv2_ok') else 'NG'})  "
            f"HV3={r.get('hv3') or ''} ({'OK' if r.get('hv3_ok') else 'NG'})  "
            f"{r.get('result') or ''}"
        )
        add_line(line, red=(r.get("result") == "FAIL"))

    return True


def _fill_salt_spray_block(table, item: Dict[str, Any]):
    start_ri = _find_section_row(table, "九．鹽霧測試")
    if start_ri is None:
        return False

    d = item.get("data") or {}
    if not d:
        return False

    for rr in range(start_ri, min(len(table.rows), start_ri + 4)):
        for c in _row_unique_cells(table.rows[rr]):
            txt = c.text or ""
            if (d.get("spec_rust_type") or "").lower() == "white":
                txt = txt.replace("□白鏽", "☑白鏽")
            if (d.get("spec_rust_type") or "").lower() == "red":
                txt = txt.replace("□紅鏽", "☑紅鏽")
            c.text = txt

    target_cell = None
    for rr in range(start_ri + 1, min(len(table.rows), start_ri + 8)):
        cells = _row_unique_cells(table.rows[rr])
        for c in cells:
            if _norm_text(c.text) == "":
                target_cell = c
                break
        if target_cell:
            break

    if not target_cell:
        return False

    target_cell.text = ""
    p = target_cell.paragraphs[0]

    def add_line(line: str, red: bool = False):
        run = p.add_run(line)
        _apply_default_run_font(run)
        if red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        run.add_break()

    rust_name = "白鏽" if (d.get("spec_rust_type") or "").lower() == "white" else "紅鏽"
    add_line(f"Hours：{rust_name} {d.get('spec_hours') or ''} H")
    add_line(f"開始時間：{d.get('start_at') or ''}")
    add_line(f"結束時間：{d.get('end_at') or ''}")
    add_line(f"實際時數：{d.get('actual_hours') or ''}", red=(d.get("result") == "FAIL"))
    add_line(
        f"結果：無鏽 {d.get('no_rust_pcs') or 0} Pcs / 白鏽 {d.get('white_rust_pcs') or 0} Pcs / 紅鏽 {d.get('red_rust_pcs') or 0} Pcs"
    )
    add_line(f"判定：{d.get('result') or ''}", red=(d.get("result") == "FAIL"))

    return True


def _fill_hydrogen_block(table, item: Dict[str, Any]):
    start_ri = _find_section_row(table, "六．氫脆測試")
    if start_ri is None:
        return False

    rows = item.get("rows", []) or []
    d = item.get("data") or {}

    target_cell = None
    for rr in range(start_ri + 1, min(len(table.rows), start_ri + 8)):
        cells = _row_unique_cells(table.rows[rr])
        for c in cells:
            if _norm_text(c.text) == "":
                target_cell = c
                break
        if target_cell:
            break

    if not target_cell:
        return False

    target_cell.text = ""
    p = target_cell.paragraphs[0]

    def add_line(line: str, red: bool = False):
        run = p.add_run(line)
        _apply_default_run_font(run)
        if red:
            run.font.color.rgb = RGBColor(255, 0, 0)
        run.add_break()

    add_line(f"Tighter torque: {d.get('tighter_torque') or ''} kg/cm")
    add_line(f"No failures after: {d.get('no_failures_after_hours') or ''} hrs.")
    add_line(
        f"OK：{item.get('pass_count') or 0}    NG：{item.get('fail_count') or 0}    結果：{item.get('result') or ''}",
        red=(item.get("result") == "FAIL"),
    )
    add_line("")

    for r in rows:
        line = (
            f"{r.get('sample_no')}. "
            f"旋入={r.get('tighten_at') or ''}  "
            f"拆卸={r.get('remove_at') or ''}  "
            f"{r.get('status') or ''}"
        )
        add_line(line, red=((r.get("status") or "").upper() == "NG"))

    return True
